import fitz  # PyMuPDF
from docx import Document
import re
import io
import os
import requests
from bs4 import BeautifulSoup

# We'll try to import textract for .doc support
try:
    import textract
except ImportError:
    textract = None

def extract_text_from_pdf(file_bytes):
    """Extracts text from a PDF file using PyMuPDF."""
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_bytes):
    """Extracts text from a Word .docx document, preserving bold markers for answer detection."""
    doc = Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        line = ""
        for run in para.runs:
            if run.bold:
                # Mark bold text with ** so regex can detect correct answers
                line += f"**{run.text}**"
            else:
                line += run.text
        if line.strip():
            full_text.append(line)
    
    # Also extract from tables (MCQs often in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)
    
    return "\n".join(full_text)

def extract_text_from_doc(file_bytes, filename):
    """Extracts text from legacy .doc files using textract."""
    if textract is None:
        return "Error: textract library not installed for .doc support."
    
    # textract usually needs a physical file, so we save it temporarily
    temp_filename = f"temp_{filename}"
    with open(temp_filename, "wb") as f:
        f.write(file_bytes)
    
    try:
        text = textract.process(temp_filename).decode('utf-8')
        return text
    except Exception as e:
        return f"Error extracting .doc: {str(e)}"
    finally:
        if os.path.exists(temp_filename):
            os.remove(temp_filename)

def parse_quiz_content(text):
    """
    Improved fallback parser using regex if AI fails.
    Detects answers from:
    - Bold markers (**Answer**)
    - Answer key lines (Answer: C, Ans: B, Correct: A)
    - Starred/marked options (*C) opt)
    """
    questions = []
    lines = text.split('\n')
    current_q = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Question pattern (e.g. "1. What is..." or "1) What is...")
        if re.match(r'^\d+[\.\)]\s+\S', line):
            if current_q and len(current_q["options"]) >= 2:
                questions.append(current_q)
            # Clean question text
            q_text = re.sub(r'\*\*', '', line)  # Remove bold markers
            current_q = {"question": q_text, "options": [], "answer": ""}
        
        # Options pattern (e.g. "A) Option" or "A. Option")
        elif re.match(r'^[A-Da-d][\.\)]\s*', line) and current_q:
            # Check if this option is marked as bold (correct answer)
            is_bold = '**' in line
            clean_option = re.sub(r'\*\*', '', line).strip()
            current_q["options"].append(clean_option)
            if is_bold and not current_q["answer"]:
                current_q["answer"] = clean_option
        
        # Answer key line (e.g. "Answer: C" or "Ans: B" or "Correct Answer: A)")
        elif re.match(r'^(Answer|Ans|Correct\s*Answer)\s*[:\-]?\s*[A-Da-d]', line, re.IGNORECASE) and current_q:
            match = re.search(r'[A-Da-d][\.\)]?', line)
            if match and current_q["options"]:
                letter = match.group().rstrip('.)').upper()
                # Find option matching this letter
                for opt in current_q["options"]:
                    if opt.upper().startswith(letter):
                        current_q["answer"] = opt
                        break
    
    # Add the last question
    if current_q and len(current_q["options"]) >= 2:
        questions.append(current_q)
    
    return questions

def extract_text_from_html(html_content):
    """Extracts clean text from HTML content."""
    soup = BeautifulSoup(html_content, "html.parser")
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.extract()
    text = soup.get_text(separator="\n")
    # Break into lines and remove leading/trailing space on each
    lines = (line.strip() for line in text.splitlines())
    # Break multi-headlines into a line each
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    # Drop blank lines
    text = '\n'.join(chunk for chunk in chunks if chunk)
    return text

def extract_text_from_url(url):
    """
    Downloads content from a URL.
    Returns a tuple: (content_type, content, text, filename)
    """
    # Handle Google Drive view/edit links to make them export links for PDF
    if "drive.google.com" in url and "/view" in url:
        # e.g. https://drive.google.com/file/d/1234567890/view
        match = re.search(r'/d/([^/]+)', url)
        if match:
            doc_id = match.group(1)
            url = f"https://drive.google.com/uc?export=download&id={doc_id}"
            
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()
    
    content_type = response.headers.get('Content-Type', '').lower()
    
    # Try to guess filename from header or url
    filename = "downloaded_file"
    if "content-disposition" in response.headers.lower():
        cd = response.headers.get("content-disposition")
        fname_match = re.search(r'filename="([^"]+)"', cd)
        if fname_match:
            filename = fname_match.group(1)
    else:
        # Fallback to URL part
        parts = url.split('/')
        if parts and "." in parts[-1]:
            filename = parts[-1].split('?')[0]
            
    text = ""
    if "application/pdf" in content_type or filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(response.content)
        content_type = "application/pdf"
    elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type or filename.lower().endswith(".docx"):
        text = extract_text_from_docx(response.content)
        content_type = "application/docx"
    elif "text/html" in content_type:
        text = extract_text_from_html(response.text)
        content_type = "text/html"
    elif "text/plain" in content_type:
        text = response.text
        content_type = "text/plain"
    else:
        # Fallback parsing attempt
        try:
            text = extract_text_from_pdf(response.content)
            content_type = "application/pdf"
        except:
            text = extract_text_from_html(response.text)
            content_type = "text/html"

    return content_type, response.content, text, filename
